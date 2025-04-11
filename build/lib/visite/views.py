from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from django.utils import timezone
from datetime import datetime
from connection.views import get_connected_user
from secretaire.models import ClSecretaire
from visiteur.models import ClVisiteur
from .models import ClVisite
from .forms import ClVisiteForm
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from django.db.models import Q
from django.core.exceptions import ValidationError


def visite_list(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    visites = ClVisite.objects.all().order_by('-ddvst', 'hvst')
    return render(request, 'visite/visite_list.html', {  'username':username,'visites': visites})


# Vue pour afficher les détails d'une visite
def visite_detail(request, id):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    visite = get_object_or_404(ClVisite, id=id)  # Récupérer une visite spécifique par ID
    return render(request, 'visite/visite_detail.html', {  'username':username,'visite': visite})

def visite_create(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    visiteur = ClVisiteur.objects.all().order_by('tnm')
    
    # Récupérer les secrétaires dont les directeurs sont associés
    directeurs = ClSecretaire.objects.filter(
        Q(directeur__isnull=False)  # Filtrer ceux qui ont un directeur
    )

    if request.method == 'POST':
        form = ClVisiteForm(request.POST)
        if form.is_valid():
            form.save()  # Sauvegarder la nouvelle visite
            return redirect('visite:visite_list')  # Rediriger vers la liste des visites
    else:
        form = ClVisiteForm()  # Afficher un formulaire vide

    return render(request, 'visite/visite_form.html', {
          'username':username,
        'form': form,
        'visiteur': visiteur,
        'directeurs': directeurs,  # Passer la liste des secrétaires avec directeurs associés
    })

def visite_update(request, id):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    try:
        visite = get_object_or_404(ClVisite, id=id)  # Récupérer la visite à modifier
    except ClVisite.DoesNotExist:
        # Si la visite n'existe pas, gérer l'erreur ici
        return render(request, 'error_page.html', {'error': 'Visite introuvable.'})
    
    # Récupérer la liste des visiteurs et des secrétaires associés à des directeurs
    visiteurs = ClVisiteur.objects.all().order_by('tnm')
    directeurs = ClSecretaire.objects.filter(directeur__isnull=False)

    if request.method == 'POST':
        form = ClVisiteForm(request.POST, instance=visite)
        if form.is_valid():
            try:
                form.save()  # Sauvegarder les modifications
                return redirect('visite:visite_list')  # Rediriger vers la liste des visites
            except ValidationError as e:
                # Si une validation échoue, capturer et afficher l'erreur
                return render(request, 'visite/visite_edit.html', {
                    'form': form,
                    'visiteurs': visiteurs,
                    'directeurs': directeurs,
                    'visite': visite,
                    'error': f"Erreur de validation : {str(e)}"  # Passer l'erreur à afficher
                })
            except Exception as e:
                # Capturer d'autres erreurs inattendues
                return render(request, 'visite/visite_edit.html', {
                    'form': form,
                    'visiteurs': visiteurs,
                    'directeurs': directeurs,
                    'visite': visite,
                    'error': f"Une erreur inattendue s'est produite : {str(e)}"
                })
    else:
        form = ClVisiteForm(instance=visite)  # Afficher le formulaire avec les données existantes

    return render(request, 'visite/visite_edit.html', {
          'username':username,
        'form': form,
        'visiteurs': visiteurs,  # Passer la liste des visiteurs
        'directeurs': directeurs,  # Passer la liste des directeurs associés
        'visite': visite,  # Passer la visite à modifier
    })

def visite_search(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    # Obtenir les paramètres de recherche
    query = request.GET.get('q', '')
    tnm = request.GET.get('tnm', '')
    tpm = request.GET.get('tpm', '')
    tsx = request.GET.get('tsx', '')
    ttvst = request.GET.get('ttvst', '')
    
    # Filtrage des visites
    visites = ClVisite.objects.all()

    # Appliquer les filtres en fonction des champs remplis
    if query:
        visites = visites.filter(tobjt__icontains=query)
    if tnm:
        visites = visites.filter(idvstr__tnm__icontains=tnm)
    if tpm:
        visites = visites.filter(idvstr__tpm__icontains=tpm)
    if tsx:
        visites = visites.filter(idvstr__tsx__icontains=tsx)
    if ttvst:
        visites = visites.filter(idvstr__ttvst__icontains=ttvst)
    
    return render(request, 'visite/visite_list.html', {  'username':username,'visites': visites, 'query': query, 'tnm': tnm, 'tpm': tpm, 'tsx': tsx, 'ttvst': ttvst})

from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.utils import timezone
from datetime import datetime
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .models import ClVisite


def visite_impression(request):
    username = get_connected_user(request)

    if not username:
        return redirect('login')

    today = timezone.now().date()
    mois = today.month
    annee = today.year

    visites = ClVisite.objects.filter(ddvst__month=mois, ddvst__year=annee)

    doc = Document()

    # Orientation paysage
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    doc.add_heading('Liste des Visites', 0)
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Visite ID'
    hdr_cells[1].text = 'Nom Visiteur'
    hdr_cells[2].text = 'Prénom Visiteur'
    hdr_cells[3].text = 'Directeur'
    hdr_cells[4].text = 'Objet'
    hdr_cells[5].text = 'Type de Visiteur'
    hdr_cells[6].text = 'Date Visite'
    hdr_cells[7].text = 'Statut'
    hdr_cells[8].text = 'Motif'

    for visite in visites:
        row_cells = table.add_row().cells
        row_cells[0].text = str(visite.id)
        row_cells[1].text = visite.idvstr.tnm
        row_cells[2].text = visite.idvstr.tpm
        row_cells[3].text = f"{visite.iddirecteur.tnm} {visite.iddirecteur.tpm}"
        row_cells[4].text = visite.tobjt
        row_cells[5].text = visite.ttvst
        row_cells[6].text = visite.ddvst.strftime('%d/%m/%Y')
        row_cells[7].text = visite.tsttvst
        row_cells[8].text = visite.tmtf if visite.tmtf else 'N/A'

    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
    run_date.bold = True
    run_date.font.size = Pt(10)

    for _ in range(3):
        doc.add_paragraph()

    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}   {' ' * 10}")
    ref_run.bold = True
    ref_run.font.size = Pt(10)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="visites.docx"'
    doc.save(response)
    return response


def generate_word(request):
    username = get_connected_user(request)

    if not username:
        return redirect('login')

    if request.method == 'POST':
        selected_ids = request.POST.getlist('visite_select[]')

        if not selected_ids:
            return HttpResponse("Aucune visite sélectionnée.", status=400)

        visites = ClVisite.objects.filter(id__in=selected_ids)
        doc = Document()

        # Orientation paysage
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        doc.add_heading("Liste des Visites", 0)
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Visite ID'
        hdr_cells[1].text = 'Nom Visiteur'
        hdr_cells[2].text = 'Prénom Visiteur'
        hdr_cells[3].text = 'Directeur'
        hdr_cells[4].text = 'Objet'
        hdr_cells[5].text = 'Type de Visiteur'
        hdr_cells[6].text = 'Date Visite'
        hdr_cells[7].text = 'Statut'
        hdr_cells[8].text = 'Motif'

        for visite in visites:
            row_cells = table.add_row().cells
            row_cells[0].text = str(visite.id)
            row_cells[1].text = visite.idvstr.tnm
            row_cells[2].text = visite.idvstr.tpm
            row_cells[3].text = f"{visite.iddirecteur.tnm} {visite.iddirecteur.tpm}"
            row_cells[4].text = visite.tobjt
            row_cells[5].text = visite.ttvst
            row_cells[6].text = visite.ddvst.strftime('%d/%m/%Y')
            row_cells[7].text = visite.tsttvst
            row_cells[8].text = visite.tmtf if visite.tmtf else 'N/A'

        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
        run_date.bold = True
        run_date.font.size = Pt(10)

        for _ in range(3):
            doc.add_paragraph()

        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}   {' ' * 10}")
        ref_run.bold = True
        ref_run.font.size = Pt(10)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        filename = f"Visites_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response

    return HttpResponse("Méthode non autorisée", status=405)
