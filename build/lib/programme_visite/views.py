import json
from pyexpat.errors import messages
from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from zipfile import ZipFile
from connection.views import get_connected_user
from secretaire.models import ClSecretaire
from visite.models import ClVisite
from django.utils import timezone
from datetime import datetime
from .models import ClProgrammeVisite
from .forms import ClProgrammeVisiteForm
from django.shortcuts import get_object_or_404
from .models import ClProgrammeVisite
from docx import Document
from django.utils.timezone import now
from django.db.models import Q
from django.shortcuts import render, get_object_or_404, redirect
from django.utils.timezone import now
from django.db.models import Q  # Pour les requêtes complexes avec OR / AND
from django.http import JsonResponse
from io import BytesIO
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from docx import Document
import json
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from django.utils import timezone
from docx.shared import Pt


# Vue pour afficher la liste des programmes de visite
def programme_visite_list(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    programmes = ClProgrammeVisite.objects.all()  # Récupérer tous les programmes de visite
    return render(request, 'programme_visite/programme_visite_list.html', {  'username':username,'programmes': programmes})

# Vue pour afficher les détails d'un programme de visite
def programme_visite_detail(request, id):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    programme = get_object_or_404(ClProgrammeVisite, id=id)  # Récupérer un programme de visite par ID
    return render(request, 'programme_visite/programme_visite_detail.html', {  'username':username,'programme': programme})


def programme_visite_create(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    today = now().date()

    # Récupérer les visites confirmées et qui ne sont pas déjà programmées
    visites = ClVisite.objects.filter(
        tsttvst='confirmé',
        ddvst__lte=today
    ).exclude(id__in=ClProgrammeVisite.objects.values_list('idvst_id', flat=True)).order_by('-ddvst')

    # Vérifier qu'il existe des visites
    if visites.exists():
        # Récupérer la secrétaire la plus récente associée au directeur de la première visite
        secretaire_recente = ClSecretaire.objects.filter(
            directeur=visites.first().iddirecteur,
            dsb__lte=now(),  # Assurez-vous que la date de début (dsb) est inférieure ou égale à la date actuelle
        ).filter(
            Q(ddf__isnull=True) | Q(ddf__lte=now())  # Filtrer si ddf est None ou inférieur à la date actuelle
        ).order_by('-dsb').first()  # Trier par dsb de manière décroissante
    else:
        secretaire_recente = None

    # Récupérer un formulaire pour la création de programme de visite
    form = ClProgrammeVisiteForm(request.POST or None)

    if request.method == 'POST' and form.is_valid():
        form.save()
        return redirect('programme_visite:programme_visite_list')

    return render(request, 'programme_visite/programme_visite_form.html', {
          'username':username,
        'form': form,
        'visites': visites,
        'secretaire_recente': secretaire_recente  # Passer la secrétaire récente au template
    })
def programme_visite_list(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    query = request.GET.get('q', '')

    if query:
        programmes = ClProgrammeVisite.objects.filter(
            Q(idvst__idvstr__tnm__icontains=query) |  # Recherche par nom du visiteur (assurez-vous que 'tnm' est un champ existant dans 'ClVisiteur')
            Q(idvst__ddvst__icontains=query) |  # Recherche dans la date de visite de 'ClVisite'
            Q(ddpvst__icontains=query) |  # Recherche dans la date programmée de visite
            Q(tsttpvst__icontains=query) |  # Recherche dans le statut
            Q(motif__icontains=query)  # Recherche dans le motif
        )
    else:
        programmes = ClProgrammeVisite.objects.all()

    return render(request, 'programme_visite/programme_visite_list.html', {  'username':username,'programmes': programmes, 'query': query})

def programme_visite_update(request, id):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    programme = get_object_or_404(ClProgrammeVisite, id=id)

    # Étape 1️⃣ : récupérer la visite liée à ce programme
    visite = programme.idvst

    # Étape 2️⃣ : récupérer le directeur lié à cette visite
    directeur = visite.iddirecteur

    # Étape 3️⃣ : récupérer la secrétaire la plus récente de ce directeur
    secretaire_recente = ClSecretaire.objects.filter(
        directeur=directeur,
        dsb__lte=now()  # Secrétaire déjà en poste
    ).filter(
          Q(ddf__isnull=True) | Q(ddf__gte=now())  # Soit toujours en poste, soit fin de poste dans le futur
    ).order_by('-dsb').first()  # La plus récente

    if request.method == 'POST':
        form = ClProgrammeVisiteForm(request.POST, instance=programme)
        if form.is_valid():
            form.save()
            return redirect('programme_visite:programme_visite_list')
    else:
        form = ClProgrammeVisiteForm(instance=programme)

        # Si pas déjà définie, on propose une secrétaire par défaut (pré-remplissage)
        if secretaire_recente and not programme.secretaire:
            form.fields['secretaire'].initial = secretaire_recente.id

    return render(request, 'programme_visite/programme_visite_edit.html', {
          'username':username,
        'form': form,
        'programme': programme,
        'visite': visite,
        'directeur': directeur,
        'secretaire_recente': secretaire_recente
    })


from io import BytesIO
from zipfile import ZipFile
from django.shortcuts import get_object_or_404, redirect
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from django.utils import timezone
from .models import ClProgrammeVisite

# Fonction pour imprimer le programme de visite
def imprimer_programme_visite(request, id):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    programme = get_object_or_404(ClProgrammeVisite, id=id)
    doc = Document()

    # Orientation paysage
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    doc.add_heading(f"Programme de Visite", level=0)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    add_row("ID Programme", str(programme.id))
    add_row("Visiteur", f"{programme.idvst.idvstr.tnm} {programme.idvst.idvstr.tpm}")
    add_row("Date de visite", programme.ddpvst.strftime('%d/%m/%Y'))
    add_row("Heure début", programme.hdbt.strftime('%H:%M'))
    if programme.hhf:
        add_row("Heure fin", programme.hhf.strftime('%H:%M'))
    add_row("Statut", programme.tsttpvst)
    if programme.motif:
        add_row("Motif", programme.motif)
    if programme.secretaire:
        add_row("Secrétaire", f"{programme.secretaire.tnm} {programme.secretaire.tpm}")
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
    run_date.bold = True
    run_date.font.size = Pt(10)

    for _ in range(3):
        doc.add_paragraph()

    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}   {' ' * 10}")
    ref_run.bold = True
    ref_run.font.size = Pt(10)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'inline; filename="programme_visite_{programme.id}.docx"'
    doc.save(response)
    return response

# Fonction pour imprimer la liste des programmes de visite
def imprimer_liste_programmes_visites(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    today = timezone.now().date()
    mois = today.month
    annee = today.year

    # Filtrer les programmes du mois en cours
    programmes = ClProgrammeVisite.objects.filter(ddpvst__month=mois, ddpvst__year=annee)

    doc = Document()

    # Orientation paysage
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    doc.add_heading("Liste des Programmes de Visite", 0)

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Visiteur'
    hdr_cells[2].text = 'Date'
    hdr_cells[3].text = 'Heure Début'
    hdr_cells[4].text = 'Heure Fin'
    hdr_cells[5].text = 'Statut'
    hdr_cells[6].text = 'Secrétaire'

    for prog in programmes:
        row_cells = table.add_row().cells
        row_cells[0].text = str(prog.id)
        row_cells[1].text = f"{prog.idvst.idvstr.tnm} {prog.idvst.idvstr.tpm}"
        row_cells[2].text = prog.ddpvst.strftime('%d/%m/%Y')
        row_cells[3].text = prog.hdbt.strftime('%H:%M')
        row_cells[4].text = prog.hhf.strftime('%H:%M') if prog.hhf else '---'
        row_cells[5].text = prog.tsttpvst
        row_cells[6].text = f"{prog.secretaire.tnm} {prog.secretaire.tpm}" if prog.secretaire else '---'

    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
    run_date.bold = True
    run_date.font.size = Pt(10)

    for _ in range(3):
        doc.add_paragraph()

    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}   {' ' * 10}")
    ref_run.bold = True
    ref_run.font.size = Pt(10)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'inline; filename="liste_programmes_visites.docx"'
    doc.save(response)
    return response

# Fonction pour générer un fichier zip contenant les programmes sélectionnés
def generate_word(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    if request.method == 'POST':
        selected_programmes = request.POST.getlist('programme_select[]')

        if not selected_programmes:
            return HttpResponse("Aucun programme sélectionné.", status=400)

        programmes = ClProgrammeVisite.objects.filter(id__in=selected_programmes)

        zip_buffer = BytesIO()
        with ZipFile(zip_buffer, 'w') as zip_file:
            for programme in programmes:
                doc = Document()

                # Orientation paysage
                section = doc.sections[-1]
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width, section.page_height = section.page_height, section.page_width

                doc.add_heading("Programme", 0)

                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Champ'
                hdr_cells[1].text = 'Valeur'

                row_cells = table.add_row().cells
                row_cells[0].text = 'Date programmée'
                row_cells[1].text = str(programme.ddpvst)

                row_cells = table.add_row().cells
                row_cells[0].text = 'Heure début'
                row_cells[1].text = str(programme.hdbt)

                row_cells = table.add_row().cells
                row_cells[0].text = 'Heure fin'
                row_cells[1].text = str(programme.hhf)

                row_cells = table.add_row().cells
                row_cells[0].text = 'Statut'
                row_cells[1].text = programme.tsttpvst

                row_cells = table.add_row().cells
                row_cells[0].text = 'Motif'
                row_cells[1].text = programme.motif

                row_cells = table.add_row().cells
                row_cells[0].text = 'Secrétaire'
                row_cells[1].text = f"{programme.secretaire.tnm} {programme.secretaire.tpm}"

                doc.add_paragraph()
                date_para = doc.add_paragraph()
                date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
                run_date.bold = True
                run_date.font.size = Pt(10)

                for _ in range(3):
                    doc.add_paragraph()

                ref_para = doc.add_paragraph()
                ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}   {' ' * 10}")
                ref_run.bold = True
                ref_run.font.size = Pt(10)
                word_buffer = BytesIO()
                doc.save(word_buffer)
                word_buffer.seek(0)
                zip_file.writestr(f"programme_{programme.idvst}.docx", word_buffer.read())

        zip_buffer.seek(0)
        response = HttpResponse(zip_buffer.read(), content_type='application/zip')
        response['Content-Disposition'] = 'attachment; filename="programmes_visites.zip"'
        return response

    return HttpResponse("Méthode non autorisée", status=405)
