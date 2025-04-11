from django.shortcuts import render, redirect
from connection.views import get_connected_user
from django.db import models
from django.db.models import Avg, Count, F
import matplotlib.pyplot as plt
import io
import base64

from programme_visite.models import ClProgrammeVisite
from secretaire.models import ClSecretaire
from visite.models import ClVisite
from visiteur.models import ClVisiteur

def tatitistique_view(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    # Récupérer toutes les visites
    visites = ClVisite.objects.all().order_by('-ddvst', 'hvst')

    # Récupérer les visites pour un directeur spécifique (en filtrant par le champ tstt, qui contient le rôle)
    directeurs_visites = visites.filter(iddirecteur__tstt='Directeur')

    # Calculer les statistiques pour l'histogramme de l'évolution des visites chaque année
    visites_par_annee = directeurs_visites.annotate(annee_visite=models.functions.TruncYear('ddvst')) \
                                          .values('annee_visite') \
                                          .annotate(nombre_visites=Count('id')) \
                                          .order_by('annee_visite')

    # Générer un graphique de l'évolution des visites par année pour un directeur
    years = [str(v['annee_visite'].year) for v in visites_par_annee]
    counts = [v['nombre_visites'] for v in visites_par_annee]

    fig, ax = plt.subplots()
    ax.bar(years, counts)
    ax.set_title("Évolution des visites par année (Directeur)")
    ax.set_xlabel("Année")
    ax.set_ylabel("Nombre de visites")

    # Convertir le graphique en image base64 pour l'intégrer dans le template
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    image_png = buffer.getvalue()
    graph_url = base64.b64encode(image_png).decode('utf-8')

    # Progression des heures moyennes d'un programme de visite
    heures_par_annee = directeurs_visites.annotate(annee_visite=models.functions.TruncYear('ddvst')) \
                                         .values('annee_visite') \
                                         .annotate(heure_moyenne=Avg('hvst')) \
                                         .order_by('annee_visite')

    # Extraire les données pour l'histogramme
    annees_heures = [str(v['annee_visite'].year) for v in heures_par_annee]
    heures_moyennes = [v['heure_moyenne'] for v in heures_par_annee]

    fig2, ax2 = plt.subplots()
    ax2.bar(annees_heures, heures_moyennes)
    ax2.set_title("Progression des heures moyennes d'un programme de visite par année")
    ax2.set_xlabel("Année")
    ax2.set_ylabel("Heure moyenne")

    # Convertir en base64
    buffer2 = io.BytesIO()
    plt.savefig(buffer2, format='png')
    buffer2.seek(0)
    image_png2 = buffer2.getvalue()
    graph_url2 = base64.b64encode(image_png2).decode('utf-8')

    # Récupérer la liste des secrétaires et leurs directeurs
    secretaires = ClSecretaire.objects.all().select_related('directeur')
    secretaires_data = [
        {'secretaire': sec.tnm+' '+ sec.tpm+' '+sec.ttvst , 'directeur': sec.directeur.tnm+'  '+sec.directeur.tpm+' '+sec.directeur.ttvst}
        for sec in secretaires
    ]

    # Récupérer la liste des visiteurs
    visiteurs = ClVisiteur.objects.all()

    # Récupérer la liste des programmes de visite confirmés et annulés
    programmes = ClProgrammeVisite.objects.all()
    programmes_confirmes = programmes.filter(tsttpvst='Confirmé')
    programmes_annules = programmes.filter(tsttpvst='Annulé')

    # Passer toutes les informations à la template
    return render(request, 'statistique.html', {
        'username': username,
        'visites': visites,
        'graph_url': graph_url,
        'graph_url2': graph_url2,
        'directeurs_visites': directeurs_visites,
        'visites_par_annee': visites_par_annee,
        'heures_par_annee': heures_par_annee,
        'secretaires_data': secretaires_data,
        'visiteurs': visiteurs,
        'programmes_confirmes': programmes_confirmes,
        'programmes_annules': programmes_annules
    })

# fichier : users/views.py (ou un fichier views.py central selon ton organisation)

from django.shortcuts import render
from visiteur.models import ClVisiteur
from directeur.models import ClDirecteur
from secretaire.models import ClSecretaire

def liste_utilisateurs(request):
    visiteurs = ClVisiteur.objects.all()
    directeurs = ClDirecteur.objects.all()
    secretaires = ClSecretaire.objects.select_related('directeur').all()  # chargement du directeur lié

    context = {
        'visiteurs': visiteurs,
        'directeurs': directeurs,
        'secretaires': secretaires,
    }
    return render(request, 'liste_utilisateurs.html', context)

# views.py

from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from visiteur.models import ClVisiteur
from directeur.models import ClDirecteur
from secretaire.models import ClSecretaire
from docx import Document
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from visiteur.models import ClVisiteur
from directeur.models import ClDirecteur
from secretaire.models import ClSecretaire
from docx import Document
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from visiteur.models import ClVisiteur
from directeur.models import ClDirecteur
from secretaire.models import ClSecretaire
from docx import Document

@csrf_exempt  # Utiliser @csrf_exempt si vous avez un problème avec le token CSRF, mais veillez à sécuriser cette route autrement
def export_word(request):
    if request.method == "POST":
        # Récupération des identifiants envoyés via POST
        visiteurs_ids = request.POST.getlist("visiteurs")
        directeurs_ids = request.POST.getlist("directeurs")
        secretaires_ids = request.POST.getlist("secretaires")

        if not visiteurs_ids and not directeurs_ids and not secretaires_ids:
            return JsonResponse({'error': 'Aucun utilisateur sélectionné.'}, status=400)

        # Récupération des objets à partir des IDs
        visiteurs = ClVisiteur.objects.filter(id__in=visiteurs_ids)
        directeurs = ClDirecteur.objects.filter(id__in=directeurs_ids)
        secretaires = ClSecretaire.objects.filter(id__in=secretaires_ids)

        # Création du document Word
        doc = Document()
        doc.add_heading('Résumé des utilisateurs sélectionnés', 0)

        # SECTION VISITEURS
        doc.add_heading('Visiteurs', level=1)
        if visiteurs.exists():
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Nom'
            hdr_cells[1].text = 'Prénom'
            hdr_cells[2].text = 'Sexe'
            hdr_cells[3].text = 'Email'
            hdr_cells[4].text = 'Téléphone'
            for v in visiteurs:
                row = table.add_row().cells
                row[0].text = v.tnm or ""
                row[1].text = v.tpm or ""
                row[2].text = v.tsx or ""
                row[3].text = v.teml or ""
                row[4].text = v.tphne or ""
        else:
            doc.add_paragraph("Aucun visiteur sélectionné.")

        # SECTION DIRECTEURS
        doc.add_heading('Directeurs', level=1)
        if directeurs.exists():
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Nom'
            hdr_cells[1].text = 'Prénom'
            hdr_cells[2].text = 'Sexe'
            hdr_cells[3].text = 'Email'
            hdr_cells[4].text = 'Téléphone'
            for d in directeurs:
                row = table.add_row().cells
                row[0].text = d.tnm or ""
                row[1].text = d.tpm or ""
                row[2].text = d.tsx or ""
                row[3].text = d.teml or ""
                row[4].text = d.tphne or ""
        else:
            doc.add_paragraph("Aucun directeur sélectionné.")

        # SECTION SECRETAIRES
        doc.add_heading('Secrétaires', level=1)
        if secretaires.exists():
            table = doc.add_table(rows=1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Nom'
            hdr_cells[1].text = 'Prénom'
            hdr_cells[2].text = 'Sexe'
            hdr_cells[3].text = 'Email'
            hdr_cells[4].text = 'Téléphone'
            hdr_cells[5].text = 'Directeur'
            for s in secretaires:
                row = table.add_row().cells
                row[0].text = s.tnm or ""
                row[1].text = s.tpm or ""
                row[2].text = s.tsx or ""
                row[3].text = s.teml or ""
                row[4].text = s.tphne or ""
                row[5].text = f"{s.directeur.tnm} {s.directeur.tpm}" if s.directeur else "Non assigné"
        else:
            doc.add_paragraph("Aucun secrétaire sélectionné.")

        # Création de la réponse pour télécharger le fichier Word
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=utilisateurs_selection.docx'
        doc.save(response)
        return response

    return JsonResponse({'error': 'Méthode non autorisée. Utilisez POST.'}, status=405)
