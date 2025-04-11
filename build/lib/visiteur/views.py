from django.shortcuts import render
import os
from django.conf import settings
# Create your views here.
from django.shortcuts import render, get_object_or_404, redirect

from connection.views import get_connected_user
from .models import ClVisiteur
from .forms import ClVisiteurForm
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from django.db.models import Q  # Importer Q pour les filtres complexes
import zipfile
import io

# Affiche la liste des visiteurs
def visiteur_list(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    visiteurs = ClVisiteur.objects.all()
    return render(request, 'visiteur/visiteur_list.html', {  'username':username,'visiteurs': visiteurs})

# Affiche les détails d'un visiteur
def visiteur_detail(request, id):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    visiteur = get_object_or_404(ClVisiteur, id=id)
    return render(request, 'visiteur/visiteur_detail.html', {  'username':username,'visiteur': visiteur})

# Permet de créer un nouveau visiteur
def visiteur_create(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    if request.method == 'POST':
        form = ClVisiteurForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('visiteur:visiteur_list')
    else:
        form = ClVisiteurForm()
    return render(request, 'visiteur/visiteur_form.html', {  'username':username,'form': form})

# Permet de modifier les informations d'un visiteur
def visiteur_update(request, id):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    visiteur = get_object_or_404(ClVisiteur, id=id)
    if request.method == 'POST':
        form = ClVisiteurForm(request.POST, instance=visiteur)
        if form.is_valid():
            form.save()
            return redirect('visiteur:visiteur_list')
    else:
        form = ClVisiteurForm(instance=visiteur)
    return render(request, 'visiteur/visiteur_edit.html', {  'username':username,'form': form, 'visiteur':visiteur})

# Permet de rechercher des visiteurs
def visiteur_search(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    query = request.GET.get('q', '')
    visiteurs = ClVisiteur.objects.filter(tnm__icontains=query)  # Exemple de filtre par nom
    return render(request, 'visiteur/visiteur_list.html', {  'username':username,'visiteurs': visiteurs, 'query': query})

def visiter_search(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    # Récupérer la valeur de recherche dans la requête GET, retirer les espaces inutiles autour
    query = request.GET.get('q', '').strip()  # 'q' est le paramètre de recherche dans l'URL

    # Initialiser la queryset avec tous les visiteurs
    visiteurs = ClVisiteur.objects.all()  # Utilisation du modèle hérité

    # Si un critère de recherche est spécifié, appliquer le filtre
    if query:
        visiteurs = visiteurs.filter(
            Q(tnm__icontains=query) |  # Recherche dans le champ 'tnm' (Nom)
            Q(tpm__icontains=query) |  # Recherche dans le champ 'tpm' (Prénom)
            Q(tsx__icontains=query) |  # Recherche dans le champ 'tsx' (Sexe)
            Q(ttvst__icontains=query)  # Recherche dans le champ 'ttvst' (Visiteur)
        )

    # Retourner la page de liste des visiteurs avec les résultats de la recherche
    return render(request, 'visiteur/visiteur_list.html', {  'username':username,'visiteurs': visiteurs, 'query': query})

def visiteur_impression(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    visiteurs = ClVisiteur.objects.all()
    doc = Document()
    doc.add_heading('Liste des Visiteurs', 0)

    for visiteur in visiteurs:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        cell1 = table.cell(0, 0)
        cell2 = table.cell(0, 1)
        cell1.width = Pt(300)
        cell2.width = Pt(100)

        cell1.paragraphs[0].add_run(f"Nom: {visiteur.tnm} {visiteur.tpm}")
        cell1.paragraphs[0].add_run(f"\nEmail: {visiteur.teml}")
        cell1.paragraphs[0].add_run(f"\nNuméro de téléphone: {visiteur.tphne}")
        cell1.paragraphs[0].add_run(f"\nAdresse: {visiteur.tads}")
        cell1.paragraphs[0].add_run(f"\nType de visiteur: {visiteur.ttvst}")
        cell1.paragraphs[0].add_run(f"\nStatut: {visiteur.tstt}")

        # Image dans MEDIA_ROOT
        try:
            if visiteur.img and os.path.exists(visiteur.img.path):
                img_path = visiteur.img.path
            else:
                img_path = os.path.join(settings.MEDIA_ROOT, 'user_images/person-1824147_1280_apFMjrC.png')

            if os.path.exists(img_path):
                cell2.paragraphs[0].clear()
                cell2.paragraphs[0].add_run().add_picture(img_path, width=Pt(80), height=Pt(100))
                cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception as e:
            print(f"Erreur chargement image : {e}")

        doc.add_paragraph()

    footer_paragraph = doc.add_paragraph()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.add_run(f"Brazzaville - {datetime.now().strftime('%Y-%m-%d')}").font.size = Pt(8)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="visiteur.docx"'
    doc.save(response)
    return response


def generate_word(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    visiteur_ids = request.POST.getlist('visiteur_select[]')
    if not visiteur_ids:
        return HttpResponse("Aucun visiteur sélectionné", status=400)

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for visiteur_id in visiteur_ids:
            try:
                visiteur = ClVisiteur.objects.get(id=visiteur_id)
            except ClVisiteur.DoesNotExist:
                continue

            doc = Document()
            doc.add_heading("Fiche du Visiteur", level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'

            cell1 = table.cell(0, 0)
            cell2 = table.cell(0, 1)
            cell1.width = Pt(300)
            cell2.width = Pt(100)

            cell1.paragraphs[0].add_run(f"Nom : {visiteur.tnm} {visiteur.tpm}")
            cell1.paragraphs[0].add_run(f"\nSexe : {visiteur.tsx}")
            cell1.paragraphs[0].add_run(f"\nTéléphone : {visiteur.tphne}")
            cell1.paragraphs[0].add_run(f"\nEmail : {visiteur.teml}")
            cell1.paragraphs[0].add_run(f"\nDate de naissance : {visiteur.dns}")
            cell1.paragraphs[0].add_run(f"\nStatut : {visiteur.tstt}")

            # Image à droite
            try:
                if visiteur.img and os.path.exists(visiteur.img.path):
                    img_path = visiteur.img.path
                else:
                    img_path = os.path.join(settings.MEDIA_ROOT, 'user_images/person-1824147_1280_apFMjrC.png')

                if os.path.exists(img_path):
                    cell2.paragraphs[0].clear()
                    cell2.paragraphs[0].add_run().add_picture(img_path, width=Pt(80), height=Pt(100))
                    cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            except Exception as e:
                print(f"Erreur image pour {visiteur.tnm} {visiteur.tpm} : {e}")

            word_buffer = io.BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)
            filename = f"Visiteur_{visiteur.tnm}_{visiteur.tpm}.docx"
            zip_file.writestr(filename, word_buffer.read())

    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename=visiteurs.zip'
    return response
