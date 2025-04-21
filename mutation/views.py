from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from connection.views import get_connected_user
from .models import CLMutation
from .forms import CMutationForm
from django.db.models import Q
from django.conf import settings
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Liste des mutations
def mutation_list(request):
    username = get_connected_user(request)

    if not username:
        return redirect('login')
    mutations = CLMutation.objects.select_related('secretaire', 'directeur').all().order_by('-dsb')
    return render(request, 'mutation/mutation_list.html', {'username':username,'mutations': mutations})

# Détail d'une mutation
def mutation_detail(request, id):
    username = get_connected_user(request)

    if not username:
        return redirect('login')
    mutation = get_object_or_404(CLMutation, pk=id)
    return render(request, 'mutation/mutation_detail.html', {'username':username,'mutation': mutation})

# Création d'une mutation
def mutation_create(request):
    username = get_connected_user(request)

    if not username:
        return redirect('login')
    if request.method == 'POST':
        form = CMutationForm(request.POST)
        if form.is_valid():
            # Vérification des doublons : même secrétaire et directeur avec date de début identique
            exist = CLMutation.objects.filter(
                secretaire=form.cleaned_data['secretaire'],
                directeur=form.cleaned_data['directeur'],
                dsb=form.cleaned_data['dsb']
            ).exists()
            if exist:
                form.add_error(None, "Cette mutation existe déjà.")
            else:
                form.save()
                return redirect('mutation:mutation_list')
    else:
        form = CMutationForm()
    return render(request, 'mutation/mutation_form.html', {'username':username,'form': form})

# Mise à jour d'une mutation
def mutation_update(request, id):
    username = get_connected_user(request)

    if not username:
        return redirect('login')
    mutation = get_object_or_404(CLMutation, pk=id)
    if request.method == 'POST':
        form = CMutationForm(request.POST, instance=mutation)
        if form.is_valid():
            form.save()
            return redirect('mutation:mutation_list')
    else:
        form = CMutationForm(instance=mutation)
    return render(request, 'mutation/mutation_edit.html', {'username':username, 'form': form})

# Recherche
def mutation_search(request):
    username = get_connected_user(request)

    if not username:
        return redirect('login')
    query = request.GET.get('q', '').strip()
    if query:
        mutations = CLMutation.objects.filter(
            Q(secretaire__tnm__icontains=query) |
            Q(secretaire__tpm__icontains=query) |
            Q(directeur__tnm__icontains=query) |
            Q(directeur__tpm__icontains=query)
        )
    else:
        mutations = CLMutation.objects.all()  # ← Important : afficher tout si rien saisi

    return render(request, 'mutation/mutation_list.html', {
        'username':username,
        'mutations': mutations,
        'query': query
    })

def mutation_impression(request):
    username = get_connected_user(request)

    if not username:
        return redirect('login')  # Redirige si l'utilisateur n'est pas connecté
    
    mutations = CLMutation.objects.select_related('secretaire', 'directeur').all()

    # Créer un document Word
    doc = Document()
    doc.add_heading('Liste des Mutations', 0)

    # Ajouter le tableau centré
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Centrer le tableau

    # Entêtes du tableau
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Secrétaire'
    hdr_cells[2].text = 'Directeur'
    hdr_cells[3].text = 'Date début'
    hdr_cells[4].text = 'Date fin'

    # Remplir le tableau avec les mutations
    for idx, mutation in enumerate(mutations, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = f"{mutation.secretaire.tnm} {mutation.secretaire.tpm}"
        row_cells[2].text = f"{mutation.directeur.tnm} {mutation.directeur.tpm}"
        row_cells[3].text = mutation.dsb.strftime('%d/%m/%Y') if mutation.dsb else ''
        row_cells[4].text = mutation.ddf.strftime('%d/%m/%Y') if mutation.ddf else ''
        for _ in range(3):  # Ajouter un espacement après chaque ligne
            doc.add_paragraph()

    # Ajouter un espace avant la date de génération
    doc.add_paragraph()

    # Ajout de la date de génération
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Aligner à droite
    run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
    run_date.bold = True
    run_date.font.size = Pt(10)

    # Espacement supplémentaire
    for _ in range(3):
        doc.add_paragraph()

    # Signature de l'utilisateur
    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Aligner à droite
    ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}   {' ' * 10}")
    ref_run.bold = True
    ref_run.font.size = Pt(10)

    # Préparer la réponse HTTP pour le téléchargement du fichier
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=mutations.docx'
    doc.save(response)
    return response

def generate_word(request):
    username = get_connected_user(request)

    if not username:
        return redirect('login')  # Redirige si l'utilisateur n'est pas connecté
    
    document = Document()
    document.add_heading('Liste des Mutations', 0)

    mutations = CLMutation.objects.select_related('secretaire', 'directeur').all()

    # Ajouter un tableau centré
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Centrer le tableau

    # Entêtes du tableau
    headers = table.rows[0].cells
    headers[0].text = "#"
    headers[1].text = "Secrétaire"
    headers[2].text = "Directeur"
    headers[3].text = "Date début"
    headers[4].text = "Date fin"

    # Remplir le tableau avec les données des mutations
    for idx, m in enumerate(mutations, start=1):
        row = table.add_row().cells
        secretaire = f"{m.secretaire.tnm} {m.secretaire.tpm}" if m.secretaire else "N/A"
        directeur = f"{m.directeur.tnm} {m.directeur.tpm}" if m.directeur else "N/A"
        date_debut = m.dsb.strftime('%d/%m/%Y') if m.dsb else ""
        date_fin = m.ddf.strftime('%d/%m/%Y') if m.ddf else "En fonction"

        row[0].text = str(idx)
        row[1].text = secretaire
        row[2].text = directeur
        row[3].text = date_debut
        row[4].text = date_fin

    # Ajouter un espace avant la date de génération
    document.add_paragraph()

    # Ajout de la date de génération
    date_para = document.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Aligner à droite
    run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
    run_date.bold = True
    run_date.font.size = Pt(10)

    # Espacement supplémentaire
    for _ in range(3):
        document.add_paragraph()

    # Signature de l'utilisateur
    ref_para = document.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Aligner à droite
    ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}   {' ' * 10}")
    ref_run.bold = True
    ref_run.font.size = Pt(10)

    # Préparer la réponse HTTP pour le téléchargement du fichier
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=mutations.docx'
    document.save(response)
    return response
