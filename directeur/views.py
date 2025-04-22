# Create your views here.
from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import io
import zipfile
import requests
from django.conf import settings
from .models import ClDirecteur
from connection.views import get_connected_user
from .forms import ClDirecteurForm
from django.utils import timezone
from django.db.models import Q  # Importer Q pour les filtres complexes

# Affiche la liste des directeurs
def directeur_list(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    directeurs = ClDirecteur.objects.all()
    return render(request, 'directeur/directeur_list.html', {
          'username':username, 
          'directeurs': directeurs})

# Affiche les détails d'un directeur spécifique
def directeur_detail(request, id):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    directeur = get_object_or_404(ClDirecteur, id=id)
    return render(request, 'directeur/directeur_detail.html', {
             'username':username, 
             'directeur': directeur})

# Permet de créer un nouveau directeur
def directeur_create(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    if request.method == 'POST':
        form = ClDirecteurForm(request.POST, request.FILES)  # Ajout de request.FILES pour gérer les fichiers
        if form.is_valid():
            form.save()
            return redirect('directeur:directeur_list')
    else:
        form = ClDirecteurForm()

    return render(request, 'directeur/directeur_form.html', {
        'username': username,
        'form': form
    })

# Permet de modifier les informations d'un directeur existant
def directeur_update(request, id):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    # Récupère la date du jour
    today = timezone.now().date()

    directeur = get_object_or_404(ClDirecteur, id=id)
    if request.method == 'POST':
        form = ClDirecteurForm(request.POST, request.FILES, instance=directeur)  # Ajout de request.FILES
        if form.is_valid():
            form.save()
            return redirect('directeur:directeur_list')
    else:
        form = ClDirecteurForm(instance=directeur)

    return render(request, 'directeur/directeur_edit.html', {
        'username': username,
        'form': form,
        'today': today,  # Passer la date actuelle
        'directeur': directeur,  # Correction de la clé dans le context (enlever les espaces)
    })

# Permet de rechercher des directeurs en fonction d'un critère
def directeur_search(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    query = request.GET.get('q', '')
    directeurs = ClDirecteur.objects.filter(tnm__icontains=query)  # Exemple de filtre par nom
    return render(request, 'directeur/directeur_list.html', {
             'username':username, 'directeurs': directeurs, 'query': query})


def directer_search(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    query = request.GET.get('q', '').strip()
    directeurs = ClDirecteur.objects.all()

    if query:
        directeurs = directeurs.filter(
            Q(tnm__icontains=query) |
            Q(tpm__icontains=query) |
            Q(tsx__icontains=query) |
            Q(tstt__icontains=query)
        )

    return render(request, 'directeur/directeur_list.html', {
             'username':username, 
             'directeurs': directeurs, 'query': query})

# 📄 Impression globale de tous les directeurs (1 fichier Word)
def directeur_impression(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    directeurs = ClDirecteur.objects.all()
    doc = Document()
    titre = doc.add_heading('Liste des Directeurs', 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for directeur in directeurs:
        try:
            if directeur.img and os.path.exists(directeur.img.path):
                img_path = directeur.img.path
                if os.path.exists(img_path):
                    para_img = doc.add_paragraph()
                    para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = para_img.add_run()
                    run_img.add_picture(img_path, width=Inches(1.0), height=Inches(1.25))
        except Exception:
            pass  # Silencieusement ignorer les erreurs liées aux images

        table = doc.add_table(rows=10, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        champs = [
            ("Nom", f"{directeur.tnm or ''} {directeur.tpm or ''}"),
            ("Sexe", directeur.tsx or ''),
            ("Date de naissance", str(directeur.dns or '')),
            ("Lieu de naissance", directeur.tlns or ''),
            ("Adresse", directeur.tads or ''),
            ("Email", directeur.teml or ''),
            ("Téléphone", directeur.tphne or ''),
            ("Date début", str(directeur.dsb or '')),
            ("Date fin", str(directeur.ddf or '')),
            ("Statut", directeur.tstt or ''),
        ]

        for i, (label, valeur) in enumerate(champs):
            table.cell(i, 0).text = f"{label} :"
            table.cell(i, 1).text = str(valeur)
        # Espacement supplémentaire
        for _ in range(3):
            doc.add_paragraph()
    # Ajouter un espace avant la date de génération
    doc.add_paragraph()

    # Ajout de la date de génération
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Aligner à droite
    run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
    run_date.bold = True
    run_date.font.size = Pt(10)

    # Espacement supplémentaire
    for _ in range(3):
        doc.add_paragraph()

    # Signature de l'utilisateur
    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Aligner à droite
    ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}   {' ' * 10}")
    ref_run.bold = True
    ref_run.font.size = Pt(10)


    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="directeurs.docx"'
    doc.save(response)
    return response

# 📁 Impression individuelle des directeurs sélectionnés (ZIP de fichiers Word)
def generate_word(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    directeur_ids = request.POST.getlist('directeur_select[]')
    if not directeur_ids:
        return HttpResponse("Aucun directeur sélectionné", status=400)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for directeur_id in directeur_ids:
            try:
                directeur = ClDirecteur.objects.get(id=directeur_id)
            except ClDirecteur.DoesNotExist:
                continue

            doc = Document()
            titre = doc.add_heading("Fiche du Directeur", 0)
            titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

            try:
                if directeur.img and directeur.img.url:
                    image_url = request.build_absolute_uri(directeur.img.url)
                    response_img = requests.get(image_url)
                    if response_img.status_code == 200:
                        temp_image = io.BytesIO(response_img.content)
                        para_img = doc.add_paragraph()
                        para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_img = para_img.add_run()
                        run_img.add_picture(temp_image, width=Inches(1.0), height=Inches(1.25))
            except Exception:
                pass  # Silencieusement ignorer les erreurs liées aux images

            table = doc.add_table(rows=10, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            champs = [
                ("Nom", f"{directeur.tnm or ''} {directeur.tpm or ''}"),
                ("Sexe", directeur.tsx or ''),
                ("Date de naissance", str(directeur.dns or '')),
                ("Lieu de naissance", directeur.tlns or ''),
                ("Adresse", directeur.tads or ''),
                ("Email", directeur.teml or ''),
                ("Téléphone", directeur.tphne or ''),
                ("Date début", str(directeur.dsb or '')),
                ("Date fin", str(directeur.ddf or '')),
                ("Statut", directeur.tstt or ''),
            ]

            for i, (label, valeur) in enumerate(champs):
                table.cell(i, 0).text = f"{label} :"
                table.cell(i, 1).text = str(valeur)

            # Espacement supplémentaire
            for _ in range(3):
                doc.add_paragraph()

            # Ajouter un espace avant la date de génération
            doc.add_paragraph()

            # Ajout de la date de génération
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Aligner à droite
            run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
            run_date.bold = True
            run_date.font.size = Pt(10)

            # Espacement supplémentaire
            for _ in range(3):
                doc.add_paragraph()

            # Signature de l'utilisateur
            ref_para = doc.add_paragraph()
            ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Aligner à droite
            ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}   {' ' * 10}")
            ref_run.bold = True
            ref_run.font.size = Pt(10)

            word_buffer = io.BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)

            filename = f"Directeur_{directeur.tnm}_{directeur.tpm}.docx"
            zip_file.writestr(filename, word_buffer.read())

    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename=directeurs.zip'
    return response
