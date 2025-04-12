from django.shortcuts import render, redirect
from connection.views import get_connected_user
from django.db import models
from django.db.models import Avg, Count, F
import matplotlib.pyplot as plt
import io
import base64
from docx.shared import Pt
from django.utils import timezone
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from programme_visite.models import ClProgrammeVisite
from secretaire.models import ClSecretaire
from visite.models import ClVisite
from visiteur.models import ClVisiteur

def tatitistique_view(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    # Récupérer toutes les visites
    visites = ClVisite.objects.all().order_by('-ddvst', 'hvst')

    # Récupérer les visites pour un directeur spécifique (en filtrant par le champ tstt, qui contient le rôle)
    directeurs_visites = visites.filter(iddirecteur__tstt='Directeur')

    # Calculer les statistiques pour l'histogramme de l'évolution des visites chaque année
    visites_par_annee = directeurs_visites.annotate(annee_visite=models.functions.TruncYear('ddvst')) \
                                          .values('annee_visite') \
                                          .annotate(nombre_visites=Count('id')) \
                                          .order_by('annee_visite')

    # Générer un graphique de l'évolution des visites par année pour un directeur
    years = [str(v['annee_visite'].year) for v in visites_par_annee]
    counts = [v['nombre_visites'] for v in visites_par_annee]

    fig, ax = plt.subplots()
    ax.bar(years, counts)
    ax.set_title("Évolution des visites par année (Directeur)")
    ax.set_xlabel("Année")
    ax.set_ylabel("Nombre de visites")

    # Convertir le graphique en image base64 pour l'intégrer dans le template
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    image_png = buffer.getvalue()
    graph_url = base64.b64encode(image_png).decode('utf-8')

    # Progression des heures moyennes d'un programme de visite
    heures_par_annee = directeurs_visites.annotate(annee_visite=models.functions.TruncYear('ddvst')) \
                                         .values('annee_visite') \
                                         .annotate(heure_moyenne=Avg('hvst')) \
                                         .order_by('annee_visite')

    # Extraire les données pour l'histogramme
    annees_heures = [str(v['annee_visite'].year) for v in heures_par_annee]
    heures_moyennes = [v['heure_moyenne'] for v in heures_par_annee]

    fig2, ax2 = plt.subplots()
    ax2.bar(annees_heures, heures_moyennes)
    ax2.set_title("Progression des heures moyennes d'un programme de visite par année")
    ax2.set_xlabel("Année")
    ax2.set_ylabel("Heure moyenne")

    # Convertir en base64
    buffer2 = io.BytesIO()
    plt.savefig(buffer2, format='png')
    buffer2.seek(0)
    image_png2 = buffer2.getvalue()
    graph_url2 = base64.b64encode(image_png2).decode('utf-8')

    # Récupérer la liste des secrétaires et leurs directeurs
    secretaires = ClSecretaire.objects.all().select_related('directeur')
    secretaires_data = [
        {'secretaire': sec.tnm+' '+ sec.tpm+' '+sec.ttvst , 'directeur': sec.directeur.tnm+'  '+sec.directeur.tpm+' '+sec.directeur.ttvst}
        for sec in secretaires
    ]

    # Récupérer la liste des visiteurs
    visiteurs = ClVisiteur.objects.all()

    # Récupérer la liste des programmes de visite confirmés et annulés
    programmes = ClProgrammeVisite.objects.all()
    programmes_confirmes = programmes.filter(tsttpvst='Confirmé')
    programmes_annules = programmes.filter(tsttpvst='Annulé')

    # Passer toutes les informations à la template
    return render(request, 'statistique.html', {
        'username': username,
        'visites': visites,
        'graph_url': graph_url,
        'graph_url2': graph_url2,
        'directeurs_visites': directeurs_visites,
        'visites_par_annee': visites_par_annee,
        'heures_par_annee': heures_par_annee,
        'secretaires_data': secretaires_data,
        'visiteurs': visiteurs,
        'programmes_confirmes': programmes_confirmes,
        'programmes_annules': programmes_annules
    })

from django.shortcuts import render
from django.http import HttpResponse
from django.views.decorators.csrf import csrf_exempt
from docx import Document
from visiteur.models import ClVisiteur
from directeur.models import ClDirecteur
from secretaire.models import ClSecretaire
from visite.models import ClVisite
from programme_visite.models import ClProgrammeVisite

def liste_utilisateurs(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    visiteurs = ClVisiteur.objects.all()
    directeurs = ClDirecteur.objects.all()
    secretaires = ClSecretaire.objects.select_related('directeur').all()
    return render(request, 'liste_utilisateurs.html', {
        'visiteurs': visiteurs,
        'directeurs': directeurs,
        'secretaires': secretaires, 
        'username':username,
    })

@csrf_exempt
def generate_word(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    if request.method == "POST":
        visiteurs_ids = request.POST.getlist("visiteurs")
        directeurs_ids = request.POST.getlist("directeurs")
        secretaires_ids = request.POST.getlist("secretaires")

        doc = Document()
        doc.add_heading("Résumé des utilisateurs sélectionnés", 0)

        # === VISITEURS ===
        if visiteurs_ids:
            doc.add_heading("Visiteurs", level=1)
            visiteurs = ClVisiteur.objects.filter(id__in=visiteurs_ids)
            for v in visiteurs:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                fields = [
                    ("Nom complet", f"{v.tnm} {v.tpm}"),
                    ("Sexe", v.tsx),
                    ("Date de naissance", str(v.dns) if v.dns else ""),
                    ("Téléphone", v.tphne),
                    ("Email", v.teml)
                ]
                for label, value in fields:
                    if value:
                        row = table.add_row().cells
                        row[0].text = label
                        row[1].text = value

                visites = ClVisite.objects.filter(idvstr=v)

                doc.add_paragraph("")

                # Programmes Confirmés
                confirmes = ClProgrammeVisite.objects.filter(idvst__in=visites, tsttpvst='confirmé')
                if confirmes.exists():
                    doc.add_heading("Programmes Confirmés", level=3)
                    tbl_conf = doc.add_table(rows=1, cols=3)
                    tbl_conf.style = 'Light Grid Accent 1'
                    hdr = tbl_conf.rows[0].cells
                    hdr[0].text = "Date"
                    hdr[1].text = "Heure"
                    hdr[2].text = "Directeur"
                    for p in confirmes:
                        directeur = p.idvst.iddirecteur
                        row = tbl_conf.add_row().cells
                        row[0].text = str(p.ddpvst)
                        row[1].text = str(p.hdbt)
                        row[2].text = f"{directeur.tnm} {directeur.tpm}"
                    doc.add_paragraph("")  # mb-2

                # Programmes Annulés
                annules = ClProgrammeVisite.objects.filter(idvst__in=visites, tsttpvst='annulé')
                if annules.exists():
                    doc.add_heading("Programmes Annulés", level=3)
                    tbl_annul = doc.add_table(rows=1, cols=3)
                    tbl_annul.style = 'Light Grid Accent 2'
                    hdr = tbl_annul.rows[0].cells
                    hdr[0].text = "Date"
                    hdr[1].text = "Heure"
                    hdr[2].text = "Motif"
                    for p in annules:
                        row = tbl_annul.add_row().cells
                        row[0].text = str(p.ddpvst)
                        row[1].text = str(p.hdbt)
                        row[2].text = p.motif or "Non spécifié"
                    doc.add_paragraph("")  # mb-2

        # === DIRECTEURS ===
        if directeurs_ids:
            doc.add_heading("Directeurs", level=1)
            directeurs = ClDirecteur.objects.filter(id__in=directeurs_ids)
            for d in directeurs:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                fields = [
                    ("Nom complet", f"{d.tnm} {d.tpm}"),
                    ("Téléphone", d.tphne),
                    ("Email", d.teml)
                ]
                for label, value in fields:
                    if value:
                        row = table.add_row().cells
                        row[0].text = label
                        row[1].text = value

                visites_dir = ClVisite.objects.filter(iddirecteur=d)
                programmes = ClProgrammeVisite.objects.filter(idvst__in=visites_dir)

                if programmes.exists():
                    doc.add_heading("Programmes Associés", level=3)
                    tbl_prog = doc.add_table(rows=1, cols=3)
                    tbl_prog.style = 'Light Grid Accent 3'
                    hdr = tbl_prog.rows[0].cells
                    hdr[0].text = "Date"
                    hdr[1].text = "Heure"
                    hdr[2].text = "Visiteur"
                    for p in programmes:
                        visiteur = p.idvst.idvstr
                        row = tbl_prog.add_row().cells
                        row[0].text = str(p.ddpvst)
                        row[1].text = str(p.hdbt)
                        row[2].text = f"{visiteur.tnm} {visiteur.tpm}"
                    doc.add_paragraph("")  # mb-2

        # === SECRETAIRES ===
        if secretaires_ids:
            doc.add_heading("Secrétaires", level=1)
            secretaires = ClSecretaire.objects.select_related('directeur').filter(id__in=secretaires_ids)
            for s in secretaires:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                fields = [
                    ("Nom complet", f"{s.tnm} {s.tpm}"),
                    ("Téléphone", s.tphne),
                    ("Email", s.teml),
                    ("Directeur", f"{s.directeur.tnm} {s.directeur.tpm}" if s.directeur else None)
                ]
                for label, value in fields:
                    if value:
                        row = table.add_row().cells
                        row[0].text = label
                        row[1].text = value

                if s.directeur:
                    visites_dir = ClVisite.objects.filter(iddirecteur=s.directeur)
                    programmes = ClProgrammeVisite.objects.filter(idvst__in=visites_dir)

                    if programmes.exists():
                        doc.add_heading("Programmes Associés", level=3)
                        tbl_prog = doc.add_table(rows=1, cols=3)
                        tbl_prog.style = 'Light Grid Accent 4'
                        hdr = tbl_prog.rows[0].cells
                        hdr[0].text = "Date"
                        hdr[1].text = "Heure"
                        hdr[2].text = "Visiteur"
                        for p in programmes:
                            visiteur = p.idvst.idvstr
                            row = tbl_prog.add_row().cells
                            row[0].text = str(p.ddpvst)
                            row[1].text = str(p.hdbt)
                            row[2].text = f"{visiteur.tnm} {visiteur.tpm}"
                        doc.add_paragraph("")  # mb-2
        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
        run_date.bold = True
        run_date.font.size = Pt(10)

        for _ in range(3):
            doc.add_paragraph()

        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}   {' ' * 10}")
        ref_run.bold = True
        ref_run.font.size = Pt(10)
        # === Export Word ===
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="acteurs_et_programmes.docx"'
        doc.save(response)
        return response

    return HttpResponse("Méthode non autorisée", status=405)
