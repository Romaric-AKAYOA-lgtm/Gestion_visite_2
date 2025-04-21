from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from django.conf import settings
from django.db.models import Q

from datetime import datetime
import zipfile
import io
import os

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from connection.views import get_connected_user
from .models import ClVisiteur
from .forms import ClVisiteurForm

# Affiche la liste des visiteurs
def visiteur_list(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    visiteurs = ClVisiteur.objects.all()
    return render(request, 'visiteur/visiteur_list.html', {  'username':username,'visiteurs': visiteurs})

# Affiche les détails d'un visiteur
def visiteur_detail(request, id):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    visiteur = get_object_or_404(ClVisiteur, id=id)
    return render(request, 'visiteur/visiteur_detail.html', {  'username':username,'visiteur': visiteur})
def visiteur_create(request):
    username = get_connected_user(request)

    if not username:
        return redirect('login')

    if request.method == 'POST':
        form = ClVisiteurForm(request.POST, request.FILES)  # Ajouter request.FILES ici
        if form.is_valid():
            form.save()
            return redirect('visiteur:visiteur_list')
    else:
        form = ClVisiteurForm()

    return render(request, 'visiteur/visiteur_form.html', {'username': username, 'form': form})

def visiteur_update(request, id):
    username = get_connected_user(request)

    if not username:
        return redirect('login')

    visiteur = get_object_or_404(ClVisiteur, id=id)
    if request.method == 'POST':
        form = ClVisiteurForm(request.POST, request.FILES, instance=visiteur)  # Ajouter request.FILES ici
        if form.is_valid():
            form.save()
            return redirect('visiteur:visiteur_list')
    else:
        form = ClVisiteurForm(instance=visiteur)

    return render(request, 'visiteur/visiteur_edit.html', {'username': username, 'form': form, 'visiteur': visiteur})

# Permet de rechercher des visiteurs
def visiteur_search(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    query = request.GET.get('q', '')
    visiteurs = ClVisiteur.objects.filter(tnm__icontains=query)  # Exemple de filtre par nom
    return render(request, 'visiteur/visiteur_list.html', {  'username':username,'visiteurs': visiteurs, 'query': query})
def visiter_search(request):
    username = get_connected_user(request)

    if not username:
        return redirect('login')

    query = request.GET.get('q', '').strip()

    # Debugging: afficher la requête
    print(f"Recherche pour : {query}")

    visiteurs = ClVisiteur.objects.all()

    if query:
        visiteurs = visiteurs.filter(
            Q(tnm__icontains=query) |
            Q(tpm__icontains=query) |
            Q(tsx__icontains=query) |
            Q(ttvst__icontains=query)
        )
        # Debugging: afficher le nombre de résultats
        print(f"Nombre de visiteurs trouvés : {visiteurs.count()}")

    return render(request, 'visiteur/visiteur_list.html', { 
        'username': username, 
        'visiteurs': visiteurs, 
        'query': query
    })

def visiteur_impression(request):
    # Vérification si l'utilisateur est connecté
    username = get_connected_user(request)
    if not username:
        return redirect('login')

    # Récupération des visiteurs
    visiteurs = ClVisiteur.objects.all()
    doc = Document()

    # Titre du document
    titre = doc.add_heading('Liste des Visiteurs', 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Boucle pour chaque visiteur
    for visiteur in visiteurs:
        # Ajout de l'image en haut et centrée
        try:
            if visiteur.img and os.path.exists(visiteur.img.path):
                img_path = visiteur.img.path
            else:
                img_path = os.path.join(settings.MEDIA_ROOT, 'user_images/person-1824147_1280.png')

            if os.path.exists(img_path):
                para_img = doc.add_paragraph()
                para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = para_img.add_run()
                run_img.add_picture(img_path, width=Inches(1.0), height=Inches(1.25))
        except Exception as e:
            print(f"Erreur image pour {visiteur.tnm} {visiteur.tpm} : {e}")

        # Création d'un tableau pour les informations du visiteur
        table = doc.add_table(rows=10, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        champs = [
            ("Nom", f"{visiteur.tnm or ''} {visiteur.tpm or ''}"),
            ("Sexe", visiteur.tsx or ''),
            ("Date de naissance", str(visiteur.dns or '')),
            ("Lieu de naissance", visiteur.tlns or ''),
            ("Adresse", visiteur.tads or ''),
            ("Email", visiteur.teml or ''),
            ("Téléphone", visiteur.tphne or ''),
            ("Date début", str(visiteur.dsb or '')),
            ("Date fin", str(visiteur.ddf or '')),
            ("Statut", visiteur.tstt or ''),
        ]

        for i, (label, valeur) in enumerate(champs):
            table.cell(i, 0).text = f"{label} :"
            table.cell(i, 1).text = str(valeur)
        # Ajouter un espace avant la date
       #doc.add_paragraph()
            # Espacement supplémentaire
        for _ in range(3):
            doc.add_paragraph()

    # Ajouter un espace avant la date
    doc.add_paragraph()

    # Ajout de la date de génération
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
    run_date.bold = True
    run_date.font.size = Pt(10)

    # Espacement supplémentaire
    for _ in range(3):
        doc.add_paragraph()

    # Signature de l'utilisateur
    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}   {' ' * 10}")
    ref_run.bold = True
    ref_run.font.size = Pt(10)

    # Réponse HTTP avec le document généré
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="visiteurs.docx"'
    doc.save(response)
    return response


import io
import zipfile
from datetime import datetime
import requests
from django.http import HttpResponse
from django.shortcuts import redirect
from django.conf import settings
from docx import Document
from docx.shared import Inches, Pt

def generate_word(request):
    username = get_connected_user(request)
    if not username:
        return redirect('login')

    visiteur_ids = request.POST.getlist('visiteur_select[]')
    if not visiteur_ids:
        return HttpResponse("Aucun visiteur sélectionné", status=400)

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for visiteur_id in visiteur_ids:
            try:
                visiteur = ClVisiteur.objects.get(id=visiteur_id)
            except ClVisiteur.DoesNotExist:
                continue

            doc = Document()
            titre = doc.add_heading("Fiche du Visiteur", 0)
            titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 📸 Ajout de l'image (téléchargée depuis l'URL)
            try:
                if visiteur.img and visiteur.img.url:
                    image_url = request.build_absolute_uri(visiteur.img.url)
                else:
                    image_url = request.build_absolute_uri(settings.MEDIA_URL + 'user_images/person-1824147_1280.png')

                response = requests.get(image_url)
                if response.status_code == 200:
                    temp_image = io.BytesIO(response.content)
                    para_img = doc.add_paragraph()
                    para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = para_img.add_run()
                    run_img.add_picture(temp_image, width=Inches(1.0), height=Inches(1.25))
                else:
                    print(f"⚠️ Échec du téléchargement de l'image : {image_url}")
            except Exception as e:
                print(f"❌ Erreur lors du téléchargement de l'image : {e}")

            # 📋 Tableau des informations
            table = doc.add_table(rows=10, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            champs = [
                ("Nom", f"{visiteur.tnm or ''} {visiteur.tpm or ''}"),
                ("Sexe", visiteur.tsx or ''),
                ("Date de naissance", visiteur.dns.strftime('%d/%m/%Y') if visiteur.dns else ''),
                ("Lieu de naissance", visiteur.tlns or ''),
                ("Adresse", visiteur.tads or ''),
                ("Email", visiteur.teml or ''),
                ("Téléphone", visiteur.tphne or ''),
                ("Date début", visiteur.dsb.strftime('%d/%m/%Y') if visiteur.dsb else ''),
                ("Date fin", visiteur.ddf.strftime('%d/%m/%Y') if visiteur.ddf else ''),
                ("Statut", visiteur.tstt or ''),
            ]

            for i, (label, valeur) in enumerate(champs):
                table.cell(i, 0).text = f"{label} :"
                table.cell(i, 1).text = str(valeur)

            # 🕘 Date + Signature
            doc.add_paragraph()
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
            run_date.bold = True
            run_date.font.size = Pt(10)

            for _ in range(3):
                doc.add_paragraph()

            ref_para = doc.add_paragraph()
            ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}   {' ' * 10}")
            except Exception:
                ref_run = ref_para.add_run("Utilisateur inconnu")
            ref_run.bold = True
            ref_run.font.size = Pt(10)

            # 💾 Enregistrement du Word
            word_buffer = io.BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)

            nom_fichier = f"Visiteur_{visiteur.tnm}_{visiteur.tpm}.docx"
            zip_file.writestr(nom_fichier, word_buffer.read())

    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename=visiteurs.zip'
    return response
