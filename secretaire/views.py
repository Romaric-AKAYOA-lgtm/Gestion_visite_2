# Django imports
from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from django.utils import timezone
from django.utils.text import slugify
from django.db.models import Q
from django.conf import settings

# Python standard library
import os
import io
import zipfile
from datetime import datetime

# Application imports
from connection.views import get_connected_user
from directeur.models import ClDirecteur
from .models import ClSecretaire
from .forms import ClSecretaireForm

# docx (python-docx) imports
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Affiche la liste des secrétaires
def secretaire_list(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    secretaire = ClSecretaire.objects.all()
    return render(request, 'secretaire/secretaire_list.html', {  'username':username,'secretaire': secretaire})

# Affiche les détails d'un secrétaire
def secretaire_detail(request, id):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    secretaire = get_object_or_404(ClSecretaire, id=id)
    return render(request, 'secretaire/secretaire_detail.html', {  'username':username,'secretaire': secretaire})

def secretaire_create(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    today = timezone.now().date()
    directeur = ClDirecteur.objects.filter(Q(ddf__isnull=True) | Q(ddf__gte=today)).order_by('-ddf')

    if request.method == 'POST':
        form = ClSecretaireForm(request.POST, request.FILES)  # Inclure request.FILES ici
        if form.is_valid():
            form.save()
            return redirect('secretaire:secretaire_list')
    else:
        form = ClSecretaireForm()

    return render(request, 'secretaire/secretaire_form.html', {
        'username': username,
        'form': form,
        'directeur': directeur,
    })

# Permet de modifier les informations d’un secrétaire existant
def secretaire_update(request, id):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    # Récupère la date du jour
    today = timezone.now().date()
    directeur = ClDirecteur.objects.filter(Q(ddf__isnull=True) | Q(ddf__gte=today)).order_by('-ddf')

    # Récupère le secrétaire à modifier
    secretaire = get_object_or_404(ClSecretaire, id=id)

    if request.method == 'POST':
        # On instancie le formulaire avec les données POST et l'instance existante
        form = ClSecretaireForm(request.POST, request.FILES, instance=secretaire)
        if form.is_valid():
            form.save()
            return redirect('secretaire:secretaire_list')
    else:
        form = ClSecretaireForm(instance=secretaire)

    return render(request, 'secretaire/secretaire_edit.html', {
          'username':username,
        'form': form,
        'directeur':directeur,
        'today': today,  # Passe la date actuelle au template
        'secretaire': secretaire,
    })

# Permet de rechercher des secrétaires
def secretaire_search(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    query = request.GET.get('q', '')
    secretaire = ClSecretaire.objects.filter(tnm__icontains=query)  # Exemple de filtre par nom
    return render(request, 'secretaire/secretaire_list.html', {  'username':username,'secretaire': secretaire, 'query': query})


def secreter_search(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    query = request.GET.get('q', '').strip()
    secretaires = ClSecretaire.objects.all()

    if query:
        secretaires = secretaires.filter(
            Q(tnm__icontains=query) |
            Q(tpm__icontains=query) |
            Q(tsx__icontains=query) |
            Q(tstt__icontains=query)
        )

    return render(request, 'secretaire/secretaire_list.html', {  'username':username,'secretaires': secretaires, 'query': query})

def generate_word(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    secretaire_ids = request.POST.getlist('secretaire_select[]')
    if not secretaire_ids:
        return HttpResponse("Aucun secrétaire sélectionné", status=400)

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for secretaire_id in secretaire_ids:
            try:
                secretaire = ClSecretaire.objects.get(id=secretaire_id)
            except ClSecretaire.DoesNotExist:
                continue

            doc = Document()
            titre = doc.add_heading('Fiche du Secrétaire', 0)
            titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Ajout de l'image
            try:
                if secretaire.img and os.path.exists(secretaire.img.path):
                    img_path = secretaire.img.path
                else:
                    img_path = os.path.join(settings.MEDIA_ROOT, 'user_images', 'person-1824147_1280_apFMjrC.png')

                if os.path.exists(img_path):
                    para_img = doc.add_paragraph()
                    para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = para_img.add_run()
                    run_img.add_picture(img_path, width=Inches(1.0), height=Inches(1.25))
                else:
                    print(f"⚠️ Image introuvable : {img_path}")

            except Exception as e:
                print(f"Erreur image pour {secretaire.tnm} {secretaire.tpm} : {e}")

            # Création du tableau d'infos
            table = doc.add_table(rows=9, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            champs = [
                ("Nom", f"{secretaire.tnm or ''} {secretaire.tpm or ''}"),
                ("Sexe", secretaire.tsx or ''),
                ("Date de naissance", str(secretaire.dns or '')),
                ("Lieu de naissance", secretaire.tlns or ''),
                ("Adresse", secretaire.tads or ''),
                ("Email", secretaire.teml or ''),
                ("Téléphone", secretaire.tphne or ''),
                ("Statut matrimonial", secretaire.tstt or ''),
                ("Fonction", secretaire.ttvst or ''),
            ]

            for i, (label, valeur) in enumerate(champs):
                table.cell(i, 0).text = f"{label} :"
                table.cell(i, 1).text = str(valeur)

            doc.add_paragraph()
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
            run_date.bold = True
            run_date.font.size = Pt(10)

            for _ in range(3):
                doc.add_paragraph()

            ref_para = doc.add_paragraph()
            ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}   {' ' * 10}")
            ref_run.bold = True
            ref_run.font.size = Pt(10)

            # Sauvegarde du document Word en mémoire
            word_buffer = io.BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)

            filename = f"Secretaire_{slugify(secretaire.tnm)}_{slugify(secretaire.tpm)}.docx"
            zip_file.writestr(filename, word_buffer.read())

    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename=secretaires.zip'
    return response

def secretaire_impression(request):
    username = get_connected_user(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('connection:login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session

    secretaires = ClSecretaire.objects.all()
    doc = Document()

    # Titre principal
    titre = doc.add_heading('Liste des Secrétaires', 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Nombre total de secrétaires : {secretaires.count()}", style='BodyText')

    for s in secretaires:
          # Titre fiche
        heading = doc.add_heading(f'FICHE DE {s.tnm.upper()} {s.tpm.upper()}', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Ajout de l'image
        try:
            if s.img and os.path.exists(s.img.path):
                img_path = s.img.path
            else:
                img_path = os.path.join(settings.MEDIA_ROOT, 'user_images', 'person-1824147_1280_apFMjrC.png')

            if os.path.exists(img_path):
                para_img = doc.add_paragraph()
                para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = para_img.add_run()
                run_img.add_picture(img_path, width=Inches(1.0), height=Inches(1.25))
            else:
                print(f"⚠️ Image introuvable : {img_path}")
        except Exception as e:
            print(f"Erreur image pour {s.tnm} {s.tpm} : {e}")


        # Tableau des informations
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        infos = [
            ("Nom", s.tnm),
            ("Post-nom", s.tpm),
            ("Sexe", s.tsx),
            ("Téléphone", s.tphne),
            ("Adresse", s.tads),
            ("Date de naissance", s.dns.strftime('%d/%m/%Y') if s.dns else ''),
            ("Lieu de naissance", s.tlns),
            ("Email", s.teml),
            ("Statut matrimonial", s.tstt),
            ("Fonction", s.ttvst),
        ]

        for label, value in infos:
            row = table.add_row().cells
            row[0].text = f"{label} :"
            row[1].text = str(value)

        doc.add_paragraph()  # Espacement entre les fiches

    # Signature et date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_date = date_para.add_run(f"Fait à Brazzaville, le {datetime.today().strftime('%d/%m/%Y')}")
    run_date.bold = True
    run_date.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ref_run = ref_para.add_run(f"{username.user.tnm.upper()} {username.user.tpm}")
    ref_run.bold = True
    ref_run.font.size = Pt(10)

    # Génération du document Word
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="liste_secretaires.docx"'
    doc.save(response)
    return response
